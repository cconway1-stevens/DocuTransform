import markdown
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def markdown_to_html(markdown_text):
    # Convert markdown text to HTML
    return markdown.markdown(markdown_text)

def add_html_to_docx(document, html_text):
    # Simple HTML to docx text converter (handles paragraphs)
    paragraphs = html_text.split('</p>')
    for paragraph in paragraphs:
        if '<p>' in paragraph:
            content = paragraph.replace('<p>', '').strip()
            if content:
                doc_para = document.add_paragraph()
                run = doc_para.add_run(content)
                run.font.size = Pt(12)  # Set font size to 12

def convert_markdown_file_to_docx(input_file, output_file):
    # Read the markdown text file
    with open(input_file, 'r', encoding='utf-8') as file:
        markdown_text = file.read()
    
    # Convert to HTML
    html_output = markdown_to_html(markdown_text)
    
    # Create a new Document
    doc = Document()
    doc.core_properties.title = "Markdown Converted Document"
    
    # Add HTML content to the document
    add_html_to_docx(doc, html_output)
    
    # Save the document
    doc.save(output_file)

# File paths
input_file = 'text.txt'
output_file = 'output.docx'

# Convert markdown to docx
convert_markdown_file_to_docx(input_file, output_file)
